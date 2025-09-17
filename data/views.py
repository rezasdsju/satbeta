import pandas as pd
import os
from django.shortcuts import render, redirect
from django.conf import settings
from django.contrib import messages
from django.http import HttpResponse
from io import BytesIO
import docx
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

def analytics_hub_view(request):
    return render(request,'data/analytics/analytics_hub.html')

def upload_dataset(request):
    if request.method == 'POST' and request.FILES.get('dataset'):
        dataset = request.FILES['dataset']
        
        # Create media directory if it doesn't exist
        if not os.path.exists(settings.MEDIA_ROOT):
            os.makedirs(settings.MEDIA_ROOT)
        
        # Save the uploaded file
        file_path = os.path.join(settings.MEDIA_ROOT, dataset.name)
        with open(file_path, 'wb+') as destination:
            for chunk in dataset.chunks():
                destination.write(chunk)
        
        # Store file path in session
        request.session['current_dataset'] = file_path
        messages.success(request, 'Dataset uploaded successfully!')
        return redirect('bivariate_analysis')
    
    return render(request, 'data/upload_dataset.html')

def bivariate_analysis(request):
    # Check if dataset is uploaded
    if 'current_dataset' not in request.session or not os.path.exists(request.session['current_dataset']):
        messages.warning(request, 'Please upload a dataset first.')
        return redirect('upload_dataset')
    
    file_path = request.session['current_dataset']
    output = None
    feature1 = None
    feature2 = None
    
    try:
        data = pd.read_csv(file_path)
        columns = data.columns.tolist()
        
        if request.method == 'POST':
            feature1 = request.POST.get('feature1')
            feature2 = request.POST.get('feature2')
            
            if feature1 not in data.columns or feature2 not in data.columns:
                messages.error(request, 'One or both features not found in the dataset.')
            else:
                # Store selected features in session for download
                request.session['bivariate_features'] = {
                    'feature1': feature1,
                    'feature2': feature2
                }
                
                # Your bivariate analysis code
                def x_list():
                    return list(data[feature1])
                
                def y_list():
                    return list(data[feature2])
                
                def column_elements():
                    list_name1 = x_list()
                    a = []
                    for i in list_name1:
                        if i not in a:
                            a.append(i)     
                    return a
                
                def column2_elements():
                    list_name2 = y_list()
                    b = []
                    for i in list_name2:
                        if i not in b:
                            b.append(i)   
                    return b
                
                def x_elements():
                    return column_elements()
                
                def y_elements():
                    return column2_elements()
                
                def ub():
                    x = x_list()
                    y = y_list()
                    y_sums_list = []
                    yy = y_elements()
                    
                    for i in yy:
                        y_sums_list_parttions = []
                        col = column_elements()
                        
                        for k in range(len(col)):
                            y_sum = 0
                            for j in range(len(x)):
                                if col[k] == x[j] and y[j] == i:
                                    y_sum += 1
                            y_sums_list_parttions.append(y_sum)
                        y_sums_list.append(y_sums_list_parttions)
                    
                    return y_sums_list
                
                def total_percent_dataframe():
                    ind = ub()
                    yy = y_elements()
                    xx = x_elements()
                    muta = []
                    
                    for j in range(len(ind)):
                        summi = sum(ind[j])
                        muta_1 = []
                        
                        for i in range(len(ind[0])):
                            percen1 = (ind[j][i] / summi) * 100 if summi > 0 else 0
                            percen1 = round(percen1, 2)
                            wow = f"{ind[j][i]}({percen1}%)"
                            muta_1.append(wow)    
                        
                        muta.append(muta_1)
                    
                    hel = []
                    
                    for i in range(len(muta[0])):
                        df_col_first = []
                        tel = []
                        
                        for j in range(len(muta)):
                            target = muta[j][i]
                            df_col_first.append(yy[j])
                            tel.append(target)
                        
                        hhh = [xx[i], tuple(tel)]
                        hel.append(hhh)
                    
                    dic = dict(hel)
                    df = pd.DataFrame(dic, index=yy)
                    df.index.name = feature2
                    df.columns = pd.MultiIndex.from_product([[feature1], df.columns])
                    
                    return df
                
                output = total_percent_dataframe()
                # Store the result in session for download
                request.session['bivariate_result'] = output.to_html(classes='table table-striped table-bordered')
                output = request.session['bivariate_result']
        
        return render(request, 'data/analytics/bivariate.html', {
            'output': output,
            'columns': columns,
            'feature1': feature1,
            'feature2': feature2
        })
    
    except Exception as e:
        messages.error(request, f'Error reading dataset: {str(e)}')
        return redirect('upload_dataset')


"""
def download_bivariate(request, format):
    # Check if we have the analysis result in session
    if 'bivariate_result' not in request.session or 'bivariate_features' not in request.session:
        messages.error(request, 'No bivariate analysis result to download. Please run the analysis first.')
        return redirect('bivariate_analysis')
    
    try:
        # Get the features and result from session
        features = request.session['bivariate_features']
        feature1 = features['feature1']
        feature2 = features['feature2']
        
        # Create a DataFrame from the HTML result
        df = pd.read_html(request.session['bivariate_result'])[0]
        
        if format == 'csv':
            # Create CSV response
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.csv"'
            
            df.to_csv(response, index=True, encoding='utf-8')
            return response
            
        elif format == 'pdf':
            # Create PDF response
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.pdf"'
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            elements = []
            
            # Add title
            styles = getSampleStyleSheet()
            title = Paragraph(f"Bivariate Analysis: {feature1} vs {feature2}", styles['Title'])
            elements.append(title)
            
            # Convert DataFrame to list of lists for PDF table
            pdf_data = [df.columns.tolist()] + df.values.tolist()
            
            # Create table
            table = Table(pdf_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(table)
            doc.build(elements)
            
            pdf = buffer.getvalue()
            buffer.close()
            response.write(pdf)
            return response
            
        elif format == 'word':
            # Create Word document response
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.docx"'
            
            doc = docx.Document()
            doc.add_heading(f'Bivariate Analysis: {feature1} vs {feature2}', 0)
            
            # Add table to Word document
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1] + 1)
            
            # Add header row
            for j, column in enumerate(df.columns):
                table.cell(0, j + 1).text = str(column)
            
            # Add index column
            for i, index in enumerate(df.index):
                table.cell(i + 1, 0).text = str(index)
            
            # Add data cells
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    table.cell(i + 1, j + 1).text = str(df.iloc[i, j])
            
            # Save document to response
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response.write(buffer.getvalue())
            buffer.close()
            
            return response
            
    except Exception as e:
        messages.error(request, f'Error generating download: {str(e)}')
        return redirect('bivariate_analysis') """
def download_bivariate(request, format):
    # Check if we have the analysis result in session
    if 'bivariate_result' not in request.session or 'bivariate_features' not in request.session:
        messages.error(request, 'No bivariate analysis result to download. Please run the analysis first.')
        return redirect('bivariate_analysis')
    
    try:
        # Get the features and result from session
        features = request.session['bivariate_features']
        feature1 = features['feature1']
        feature2 = features['feature2']
        
        # Get the original data to create proper DataFrame
        file_path = request.session['current_dataset']
        data = pd.read_csv(file_path)
        
        # Recreate the bivariate analysis
        def perform_bivariate_analysis(data, feature1, feature2):
            def x_list():
                return list(data[feature1])
            
            def y_list():
                return list(data[feature2])
            
            def column_elements():
                list_name1 = x_list()
                a = []
                for i in list_name1:
                    if i not in a:
                        a.append(i)     
                return a
            
            def column2_elements():
                list_name2 = y_list()
                b = []
                for i in list_name2:
                    if i not in b:
                        b.append(i)   
                return b
            
            def x_elements():
                return column_elements()
            
            def y_elements():
                return column2_elements()
            
            def ub():
                x = x_list()
                y = y_list()
                y_sums_list = []
                yy = y_elements()
                
                for i in yy:
                    y_sums_list_parttions = []
                    col = column_elements()
                    
                    for k in range(len(col)):
                        y_sum = 0
                        for j in range(len(x)):
                            if col[k] == x[j] and y[j] == i:
                                y_sum += 1
                        y_sums_list_parttions.append(y_sum)
                    y_sums_list.append(y_sums_list_parttions)
                
                return y_sums_list
            
            def total_percent_dataframe():
                ind = ub()
                yy = y_elements()
                xx = x_elements()
                muta = []
                
                for j in range(len(ind)):
                    summi = sum(ind[j])
                    muta_1 = []
                    
                    for i in range(len(ind[0])):
                        percen1 = (ind[j][i] / summi) * 100 if summi > 0 else 0
                        percen1 = round(percen1, 2)
                        wow = f"{ind[j][i]}({percen1}%)"
                        muta_1.append(wow)    
                    
                    muta.append(muta_1)
                
                hel = []
                
                for i in range(len(muta[0])):
                    df_col_first = []
                    tel = []
                    
                    for j in range(len(muta)):
                        target = muta[j][i]
                        df_col_first.append(yy[j])
                        tel.append(target)
                    
                    hhh = [xx[i], tuple(tel)]
                    hel.append(hhh)
                
                dic = dict(hel)
                df = pd.DataFrame(dic, index=yy)
                df.index.name = feature2
                df.columns = pd.MultiIndex.from_product([[feature1], df.columns])
                
                return df
            
            return total_percent_dataframe()
        
        # Perform the analysis again to get proper DataFrame
        df = perform_bivariate_analysis(data, feature1, feature2)
        
        if format == 'csv':
            # Create CSV response
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.csv"'
            
            df.to_csv(response, index=True, encoding='utf-8')
            return response
            
        elif format == 'pdf':
            # Create PDF response
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.pdf"'
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            elements = []
            
            # Add title
            styles = getSampleStyleSheet()
            title = Paragraph(f"Bivariate Analysis: {feature1} vs {feature2}", styles['Title'])
            elements.append(title)
            
            # Add some space
            elements.append(Paragraph("<br/><br/>", styles['Normal']))
            
            # Convert DataFrame to list of lists for PDF table
            # First, prepare column headers
            column_headers = [df.index.name or 'Index']
            for col in df.columns:
                if isinstance(col, tuple):
                    column_headers.append(str(col[1]))
                else:
                    column_headers.append(str(col))
            
            pdf_data = [column_headers]
            
            # Add data rows
            for index, row in df.iterrows():
                row_data = [str(index)]
                for value in row:
                    row_data.append(str(value))
                pdf_data.append(row_data)
            
            # Create table
            table = Table(pdf_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4e73df')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8f9fc')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dddfeb'))
            ]))
            
            elements.append(table)
            doc.build(elements)
            
            pdf = buffer.getvalue()
            buffer.close()
            response.write(pdf)
            return response
            
        elif format == 'word':
            # Create Word document response
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="bivariate_analysis_{feature1}_vs_{feature2}.docx"'
            
            doc = docx.Document()
            doc.add_heading(f'Bivariate Analysis: {feature1} vs {feature2}', 0)
            
            # Add a paragraph with timestamp
            from datetime import datetime
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # Add table to Word document
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1] + 1)
            
            # Add header row - first cell is for index name
            table.cell(0, 0).text = df.index.name or 'Index'
            for j, column in enumerate(df.columns):
                if isinstance(column, tuple):
                    table.cell(0, j + 1).text = str(column[1])
                else:
                    table.cell(0, j + 1).text = str(column)
            
            # Add index values and data cells
            for i, index_value in enumerate(df.index):
                table.cell(i + 1, 0).text = str(index_value)
                for j, value in enumerate(df.iloc[i]):
                    table.cell(i + 1, j + 1).text = str(value)
            
            # Apply some basic styling to the table
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # Center alignment
            
            # Save document to response
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response.write(buffer.getvalue())
            buffer.close()
            
            return response
            
    except Exception as e:
        messages.error(request, f'Error generating download: {str(e)}')
        return redirect('bivariate_analysis')